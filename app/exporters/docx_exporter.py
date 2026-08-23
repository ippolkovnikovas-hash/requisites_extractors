"""
Заполнение шаблона shablon.docx значениями из RequisitesData.
Плейсхолдеры в шаблоне обёрнуты в одинарные кавычки: 'FULL_ORG_NAME'.
"""

import io
from pathlib import Path

from docx import Document

from app.schemas.requisites import RequisitesData


def fill_template(
    template_path: Path,
    requisites: RequisitesData,
    out_path: Path,
) -> Path:
    """Заполняет шаблон и сохраняет результат в файл."""
    doc = _fill(template_path, requisites)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    return out_path


def fill_template_to_bytes(
    template_path: Path,
    requisites: RequisitesData,
) -> io.BytesIO:
    """
    Заполняет шаблон и возвращает результат в памяти.

    Используется в `/generate`: готовый документ отдаётся пользователю сразу и
    не должен оставаться на диске или в репозитории (CLAUDE.md).
    """
    doc = _fill(template_path, requisites)
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def _fill(template_path: Path, requisites: RequisitesData) -> Document:
    """
    Открывает шаблон и заменяет все плейсхолдеры вида 'ALIAS' на значения из
    RequisitesData.
    """
    doc = Document(str(template_path))
    substitutions = {
        f"'{alias}'": value
        for alias, value in requisites.to_template_dict().items()
    }

    # Таблицы
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                _replace_in_cell(cell, substitutions)

    # Параграфы вне таблиц (на случай если в шаблоне есть текст вне таблицы)
    for para in doc.paragraphs:
        _replace_in_para(para, substitutions)

    return doc


def _replace_in_cell(cell, substitutions: dict[str, str]) -> None:
    for para in cell.paragraphs:
        _replace_in_para(para, substitutions)


def _replace_in_para(para, substitutions: dict[str, str]) -> None:
    """
    Замена с учётом того, что плейсхолдер может быть разбит по runs.
    Собираем полный текст параграфа, делаем замену, записываем в первый run.
    """
    full_text = "".join(run.text for run in para.runs)
    new_text = full_text
    for placeholder, value in substitutions.items():
        new_text = new_text.replace(placeholder, value)

    if new_text != full_text:
        # Кладём весь текст в первый run, остальные обнуляем
        if para.runs:
            para.runs[0].text = new_text
            for run in para.runs[1:]:
                run.text = ""
