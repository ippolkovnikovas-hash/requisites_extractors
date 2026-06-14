"""
Извлечение текста из .docx файлов через python-docx.

Собирает текст из:
  - параграфов документа
  - всех таблиц (построчно, ячейка за ячейкой)
  - колонтитулов (header / footer) каждой секции
"""

from pathlib import Path
from docx import Document
from docx.oxml.ns import qn
from loguru import logger

from app.schemas.extraction import TextExtractionResult
from app.core.enums import ExtractorType
from app.core.exceptions import TextExtractionError


def extract_docx(file_path: Path) -> TextExtractionResult:
    """
    Принимает путь к .docx, возвращает TextExtractionResult.
    Никогда не бросает исключение наружу — оборачивает в TextExtractionError.
    """
    warnings: list[str] = []

    try:
        doc = Document(str(file_path))
    except Exception as e:
        raise TextExtractionError(f"Cannot open DOCX: {e}", {"path": str(file_path)})

    parts: list[str] = []

    # --- Параграфы документа ---
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    # --- Таблицы ---
    for table_idx, table in enumerate(doc.tables):
        table_parts: list[str] = []
        for row in table.rows:
            row_cells = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    row_cells.append(cell_text)
            if row_cells:
                table_parts.append(" | ".join(row_cells))
        if table_parts:
            parts.append("\n".join(table_parts))

    # --- Колонтитулы ---
    for section in doc.sections:
        for hf in (section.header, section.footer):
            if hf is not None:
                for para in hf.paragraphs:
                    text = para.text.strip()
                    if text:
                        parts.append(text)

    full_text = "\n".join(parts)

    if not full_text.strip():
        warnings.append("DOCX extracted text is empty — document may be blank or image-only")
        logger.warning("Empty text from DOCX", path=str(file_path))

    logger.debug(
        "DOCX extraction done",
        path=file_path.name,
        chars=len(full_text),
        paragraphs=len(doc.paragraphs),
        tables=len(doc.tables),
    )

    return TextExtractionResult(
        text=full_text,
        extractor_used=ExtractorType.PYTHON_DOCX,
        pages=None,       # DOCX не имеет фиксированных страниц
        ocr_used=False,
        warnings=warnings,
    )