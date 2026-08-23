"""
Тесты REST-скачивания результатов.

Раз pipeline больше не сохраняет артефакты на диск, `/api/download` обязан
собирать запрошенный формат на лету из результата, оставшегося в памяти.
Побочный эффект — файлы не накапливаются в exports/ и не переживают ответ,
как и требует CLAUDE.md.

Реальные OCR/LLM/сеть не вызываются: run_pipeline подменяется на фиксированный
результат.
"""
import io

import pytest
from docx import Document

from app.schemas.requisites import RequisitesData
from app.schemas.validation import PipelineResult, ValidationReport

_DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


@pytest.fixture
def app():
    from app.main import create_app

    flask_app = create_app()
    flask_app.config["TESTING"] = True
    return flask_app


@pytest.fixture
def client(app):
    return app.test_client()


@pytest.fixture
def fake_result():
    return PipelineResult(
        document_id="apitest1",
        original_filename="sample.pdf",
        data=RequisitesData(
            company_name="ООО Тестовая Компания",
            inn="7744012347",
            kpp="774401001",
        ),
        validation=ValidationReport(),
        needs_review=False,
        warnings=[],
        status="done",
        fill_rate=0.19,
        raw_text_path=None,
        json_path=None,
        xlsx_path=None,
        docx_path=None,
        processing_meta={},
    )


@pytest.fixture
def docx_template(tmp_path, monkeypatch):
    doc = Document()
    doc.add_paragraph("'FULL_ORG_NAME'")
    table = doc.add_table(rows=1, cols=1)
    table.rows[0].cells[0].text = "'INN'"
    doc.save(str(tmp_path / "shablon.docx"))
    monkeypatch.chdir(tmp_path)
    return tmp_path


@pytest.fixture
def uploaded(client, monkeypatch, fake_result):
    import app.api.routes_upload as routes

    monkeypatch.setattr(
        routes, "run_pipeline", lambda path, original_filename: fake_result
    )
    response = client.post(
        "/api/extract",
        data={"file": (io.BytesIO(b"fake pdf bytes"), "sample.pdf")},
        content_type="multipart/form-data",
    )
    assert response.status_code == 200
    return response.get_json()


# ── /api/extract ─────────────────────────────────────────────────────────────


def test_extract_returns_data_without_file_paths(uploaded):
    assert uploaded["document_id"] == "apitest1"
    assert uploaded["data"]["inn"] == "7744012347"
    assert uploaded["json_path"] is None
    assert uploaded["xlsx_path"] is None
    assert uploaded["docx_path"] is None


def test_extract_leaves_nothing_in_exports(uploaded, tmp_path):
    exports = tmp_path / "exports"
    assert not exports.exists() or not list(exports.glob("*"))


# ── /api/download ────────────────────────────────────────────────────────────


def test_download_json_generated_on_the_fly(client, uploaded):
    response = client.get("/api/download/apitest1/json")
    assert response.status_code == 200
    assert "attachment" in response.headers.get("Content-Disposition", "")

    import json

    payload = json.loads(response.data.decode("utf-8"))
    assert payload["document_id"] == "apitest1"
    assert payload["data"]["inn"] == "7744012347"


def test_download_xlsx_generated_on_the_fly(client, uploaded):
    response = client.get("/api/download/apitest1/xlsx")
    assert response.status_code == 200
    assert "attachment" in response.headers.get("Content-Disposition", "")
    # XLSX — это zip-контейнер, он начинается с сигнатуры PK
    assert response.data[:2] == b"PK"


def test_download_docx_generated_on_the_fly(client, uploaded, docx_template):
    response = client.get("/api/download/apitest1/docx")
    assert response.status_code == 200
    assert response.headers["Content-Type"].startswith(_DOCX_MIME)

    doc = Document(io.BytesIO(response.data))
    paragraph_text = "\n".join(p.text for p in doc.paragraphs)
    assert "ООО Тестовая Компания" in paragraph_text


def test_download_docx_without_template_is_not_500(client, uploaded, tmp_path, monkeypatch):
    monkeypatch.chdir(tmp_path)  # shablon.docx намеренно отсутствует
    response = client.get("/api/download/apitest1/docx")
    assert response.status_code != 500
    assert "attachment" not in response.headers.get("Content-Disposition", "")


def test_download_does_not_persist_anything(client, uploaded, docx_template):
    for fmt in ("json", "xlsx", "docx"):
        assert client.get(f"/api/download/apitest1/{fmt}").status_code == 200

    exports = docx_template / "exports"
    assert not exports.exists() or not list(exports.glob("*"))


def test_download_unknown_format_is_400(client, uploaded):
    response = client.get("/api/download/apitest1/pdf")
    assert response.status_code == 400


def test_download_unknown_document_is_404(client):
    response = client.get("/api/download/doesnotexist/json")
    assert response.status_code == 404
