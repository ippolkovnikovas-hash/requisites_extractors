"""
Unit/integration-тесты app/web/routes.py: _build_review_rows, POST /upload
(с подменённым pipeline), POST /generate.

Изоляция: Flask test_client через app.web.create_app() с UPLOAD_FOLDER/
EXPORT_FOLDER, переопределёнными на подпапки tmp_path через app.config (не
через .env — .env не читается и не проверяется этими тестами вообще, конфиг
переопределяется напрямую). Реальный OCR/LLM/Ollama/сеть не вызываются —
run_pipeline подменяется monkeypatch там, где нужен upload(). Реальный
shablon.docx проекта не читается и не изменяется — используется мини-шаблон,
созданный в tmp_path, с реальным синтаксисом плейсхолдеров fill_template()
(текст в одинарных кавычках вида 'FULL_ORG_NAME'/'INN', а не {{...}}).
"""

import io
from pathlib import Path

import pytest
from docx import Document

from app.schemas.requisites import RequisitesData
from app.schemas.validation import FieldValidation, PipelineResult, ValidationReport
from app.web.routes import _build_review_rows

# ── Фикстуры ──────────────────────────────────────────────────────────────────


@pytest.fixture
def app(tmp_path):
    from app.web import create_app

    flask_app = create_app()
    flask_app.config["TESTING"] = True
    flask_app.config["UPLOAD_FOLDER"] = str(tmp_path / "uploads")
    flask_app.config["EXPORT_FOLDER"] = str(tmp_path / "exports")
    return flask_app


@pytest.fixture
def client(app):
    return app.test_client()


@pytest.fixture
def docx_template(tmp_path, monkeypatch):
    """Мини-шаблон в tmp_path: параграф с 'FULL_ORG_NAME', таблица с 'INN'."""
    doc = Document()
    doc.add_paragraph("'FULL_ORG_NAME'")
    table = doc.add_table(rows=1, cols=1)
    table.rows[0].cells[0].text = "'INN'"
    doc.save(str(tmp_path / "shablon.docx"))
    monkeypatch.chdir(tmp_path)
    return tmp_path


def _valid_form_data(**overrides):
    data = {
        "document_id": "test1234",
        "company_name": "ООО Тестовая Компания",
        "short_name": "",
        "legal_address": "",
        "postal_address": "",
        "ogrn": "1027700123450",
        "inn": "7744012347",
        "kpp": "774401001",
        "bank_name": "",
        "checking_account": "40702810200000012345",
        "correspondent_account": "30101810400000000225",
        "bik": "044525225",
        "ceo_position": "",
        "ceo_fio_full": "",
        "ceo_fio": "",
        "phone": "",
        "email": "",
    }
    data.update(overrides)
    return data


# ── _build_review_rows ───────────────────────────────────────────────────────


def test_build_review_rows_invalid_field_uses_raw_value_not_cleared_data():
    data = RequisitesData(bik=None)  # уже обнулено validate_requisites
    field_results = {
        "bik": FieldValidation(
            valid=False, raw_value="124525225", reason="wrong length: 5, expected 9"
        ),
    }
    rows = _build_review_rows(data, field_results, None)
    row = next(r for r in rows if r["key"] == "bik")
    assert row["value"] == "124525225"
    assert row["error"] is not None


def test_build_review_rows_valid_with_warning_shows_warning_not_error():
    data = RequisitesData(checking_account="40702810000000012345")
    field_results = {
        "checking_account": FieldValidation(
            valid=True,
            raw_value="40702810000000012345",
            normalized_value="40702810000000012345",
            warning="контрольный ключ счёта не совпадает с БИК",
        ),
    }
    rows = _build_review_rows(data, field_results, None)
    row = next(r for r in rows if r["key"] == "checking_account")
    assert row["error"] is None
    assert row["warning"] is not None


def test_build_review_rows_ogrn_label_13_digits():
    data = RequisitesData(ogrn="1027700123450")
    field_results = {
        "ogrn": FieldValidation(
            valid=True, raw_value="1027700123450", normalized_value="1027700123450"
        )
    }
    rows = _build_review_rows(data, field_results, None)
    row = next(r for r in rows if r["key"] == "ogrn")
    assert row["label"] == "ОГРН"


def test_build_review_rows_ogrn_label_15_digits():
    data = RequisitesData(ogrn="304500116000157")
    field_results = {
        "ogrn": FieldValidation(
            valid=True, raw_value="304500116000157", normalized_value="304500116000157"
        )
    }
    rows = _build_review_rows(data, field_results, None)
    row = next(r for r in rows if r["key"] == "ogrn")
    assert row["label"] == "ОГРНИП"


def test_build_review_rows_ogrn_label_missing_generic():
    data = RequisitesData(ogrn=None)
    field_results = {
        "ogrn": FieldValidation(
            valid=False, is_missing=True, raw_value=None, reason="field is null"
        )
    }
    rows = _build_review_rows(data, field_results, None)
    row = next(r for r in rows if r["key"] == "ogrn")
    assert row["label"] in ("ОГРН", "ОГРН/ОГРНИП")


def test_build_review_rows_source_llm_and_regex_displayed():
    data = RequisitesData(inn="7744012347", bik="044525225", company_name="ООО Тест")
    source_map = {"inn": "llm", "bik": "regex"}
    rows = _build_review_rows(data, {}, source_map)
    inn_row = next(r for r in rows if r["key"] == "inn")
    bik_row = next(r for r in rows if r["key"] == "bik")
    company_row = next(r for r in rows if r["key"] == "company_name")
    assert inn_row["source"] == "llm"
    assert bik_row["source"] == "regex"
    assert company_row["source"] is None


def test_build_review_rows_all_sixteen_fields_with_russian_labels():
    data = RequisitesData()
    rows = _build_review_rows(data, {}, None)
    assert len(rows) == 16
    keys = {r["key"] for r in rows}
    assert keys == set(RequisitesData.model_fields.keys())

    exact_labels = {
        "inn": "ИНН",
        "kpp": "КПП",
        "bik": "БИК",
        "checking_account": "Расчётный счёт",
        "correspondent_account": "Корреспондентский счёт",
    }
    by_key = {r["key"]: r for r in rows}
    for key, expected_label in exact_labels.items():
        assert by_key[key]["label"] == expected_label

    for row in rows:
        assert isinstance(row["label"], str) and row["label"].strip()
        assert any(
            "а" <= ch.lower() <= "я" or ch.lower() == "ё" for ch in row["label"]
        ), f"label for {row['key']} does not look Russian: {row['label']!r}"


# ── upload() с подменённым pipeline ──────────────────────────────────────────


def test_upload_renders_review_form_using_mocked_pipeline(client, monkeypatch):
    import app.web.routes as routes

    fake_result = PipelineResult(
        document_id="abc12345",
        original_filename="test.pdf",
        data=RequisitesData(company_name="ООО Тест", inn="7744012347"),
        validation=ValidationReport(),
        needs_review=False,
        warnings=[],
        status="done",
        fill_rate=0.5,
        raw_text_path=None,
        json_path=None,
        xlsx_path=None,
        docx_path=None,
        processing_meta={},
    )
    monkeypatch.setattr(
        routes, "run_pipeline", lambda file_path, original_filename: fake_result
    )

    response = client.post(
        "/upload",
        data={"file": (io.BytesIO(b"fake pdf bytes"), "test.pdf")},
        content_type="multipart/form-data",
    )
    body = response.get_data(as_text=True)
    assert "ООО Тест" in body
    assert "7744012347" in body


def test_upload_removes_saved_file_after_successful_processing(
    client, app, monkeypatch
):
    """
    Регрессия: загруженный документ сохранялся в UPLOAD_FOLDER и никогда не
    удалялся — реквизиты переживали обработку на диске, хотя CLAUDE.md этого
    не допускает. API-ветка (`routes_upload.py`) устроена правильно, через
    NamedTemporaryFile + unlink; веб-ветка была исключением.
    """
    import app.web.routes as routes

    fake_result = PipelineResult(
        document_id="abc12345",
        original_filename="test.pdf",
        data=RequisitesData(),
        validation=ValidationReport(),
        needs_review=False,
        warnings=[],
        status="done",
        fill_rate=0.0,
        processing_meta={},
    )
    monkeypatch.setattr(
        routes, "run_pipeline", lambda file_path, original_filename: fake_result
    )

    client.post(
        "/upload",
        data={"file": (io.BytesIO(b"fake pdf bytes"), "test.pdf")},
        content_type="multipart/form-data",
    )

    upload_folder = Path(app.config["UPLOAD_FOLDER"])
    assert list(upload_folder.iterdir()) == []


def test_upload_removes_saved_file_even_when_pipeline_raises(client, app, monkeypatch):
    """Файл не должен остаться на диске и в случае, когда обработка упала —
    иначе каждая ошибка pipeline оставляет за собой реквизиты в uploads/."""
    import app.web.routes as routes

    def boom(file_path, original_filename):
        raise RuntimeError("pipeline boom")

    monkeypatch.setattr(routes, "run_pipeline", boom)

    client.post(
        "/upload",
        data={"file": (io.BytesIO(b"fake pdf bytes"), "test.pdf")},
        content_type="multipart/form-data",
    )

    upload_folder = Path(app.config["UPLOAD_FOLDER"])
    assert list(upload_folder.iterdir()) == []


# ── POST /generate ────────────────────────────────────────────────────────────


def test_generate_with_valid_edited_fields_returns_docx_with_paragraph_and_table_substitution(
    client, docx_template
):
    response = client.post("/generate", data=_valid_form_data())
    assert response.status_code == 200
    assert "attachment" in response.headers.get("Content-Disposition", "")

    doc = Document(io.BytesIO(response.data))
    paragraph_text = "\n".join(p.text for p in doc.paragraphs)
    table_text = "\n".join(
        cell.text for table in doc.tables for row in table.rows for cell in row.cells
    )
    assert "ООО Тестовая Компания" in paragraph_text
    assert "7744012347" in table_text
    assert "FULL_ORG_NAME" not in paragraph_text
    assert "'INN'" not in table_text


def test_generate_blocked_without_confirm_invalid_keeps_form_values(
    client, docx_template
):
    response = client.post("/generate", data=_valid_form_data(inn="1234567890"))
    assert response.status_code == 422
    assert "attachment" not in response.headers.get("Content-Disposition", "")
    body = response.get_data(as_text=True)
    assert "1234567890" in body
    assert not (docx_template / "exports").exists() or not list(
        (docx_template / "exports").glob("*.docx")
    )


def test_generate_with_confirm_invalid_allows_docx_with_invalid_value(
    client, docx_template
):
    response = client.post(
        "/generate",
        data=_valid_form_data(inn="1234567890", confirm_invalid="on"),
    )
    assert response.status_code == 200
    assert "attachment" in response.headers.get("Content-Disposition", "")
    doc = Document(io.BytesIO(response.data))
    table_text = "\n".join(
        cell.text for table in doc.tables for row in table.rows for cell in row.cells
    )
    assert "1234567890" in table_text


def test_generate_empty_optional_field_does_not_block(client, docx_template):
    response = client.post("/generate", data=_valid_form_data(email=""))
    assert response.status_code == 200
    assert "attachment" in response.headers.get("Content-Disposition", "")


def test_generate_missing_required_field_is_not_same_as_invalid_and_does_not_block(
    client, docx_template
):
    """Пустой bik (не введён) — это не то же самое, что введённый неверный bik,
    и не требует confirm_invalid."""
    response = client.post("/generate", data=_valid_form_data(bik=""))
    assert response.status_code == 200
    assert "attachment" in response.headers.get("Content-Disposition", "")


def test_generate_blocked_by_invalid_email_without_confirm(client, docx_template):
    response = client.post("/generate", data=_valid_form_data(email="not-an-email"))
    assert response.status_code == 422
    assert "attachment" not in response.headers.get("Content-Disposition", "")
    body = response.get_data(as_text=True)
    assert "not-an-email" in body


def test_generate_missing_shablon_returns_safe_error(client, tmp_path, monkeypatch):
    monkeypatch.chdir(tmp_path)  # shablon.docx намеренно не создаём
    response = client.post("/generate", data=_valid_form_data())
    assert response.status_code != 500
    assert "attachment" not in response.headers.get("Content-Disposition", "")


def test_generate_does_not_persist_files_in_export_folder(client, docx_template):
    response = client.post("/generate", data=_valid_form_data())
    assert response.status_code == 200
    exports_dir = docx_template / "exports"
    assert not exports_dir.exists() or not list(exports_dir.glob("*"))


def test_generate_without_document_id_generates_safe_filename(client, docx_template):
    data = _valid_form_data()
    del data["document_id"]
    response = client.post("/generate", data=data)
    assert response.status_code == 200
    content_disposition = response.headers.get("Content-Disposition", "")
    assert "attachment" in content_disposition
    assert ".docx" in content_disposition
    exports_dir = docx_template / "exports"
    assert not exports_dir.exists() or not list(exports_dir.glob("*"))
